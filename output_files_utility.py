from docx import Document
from docx.shared import Inches
from datetime import datetime

class GenerateReport():
    """

    """

    def __init__(self, file_name):
        """

        :param file_name:
        """
        self.file_name =  file_name
        self.document = Document()

    def add_sample_info(self, FCS_data, hx = False):
        """

        :param run_date:
        :param sample_date:
        :param mrn:
        :param patient:
        :param name:
        :return:
        """
        #print()
        if hx:
            self.document.add_paragraph(" ".join(["run date", str(FCS_data["rundate"]), "sample date",
                                                  str(FCS_data["sampledate"]),"MRN", str(FCS_data["sampleMRN"])]))
        else:
            self.document.add_paragraph(" ".join(["run date", str(FCS_data["rundate"]), "sample date",
                                                  str(FCS_data["sampledate"]), "MRN", str(FCS_data["sampleMRN"])]),
                                        style="Intense Quote")



    def add_delta_check_info(self, info_dict):
        """

        :param info_dict:
        :return:
        """
        for k, v in info_dict.items():

            self.document.add_paragraph(str(k) + ":" + str(v))
    def add_as_newsample(self):
        """

        :return:
        """
        self.document.add_paragraph("NEW SAMPLE")

    def add_title(self):
        """

        :return:
        """
        pass

        datetime_stamp = datetime.now()
        str_timestample = datetime_stamp.strftime("%d-%b-%Y-%H:%M")
        self.document.add_paragraph("This report was generated on " + str_timestample, style = "Title")

    def add_graph(self, graph_link):
        """

        :param graph_link:
        :return:

        """

        self.document.add_picture (graph_link, width = Inches(7))



    def save_document(self, output_folder):
        """

        :param output_folder:
        :param output_name:
        :return:
        """
        self.document.save(output_folder +"/" +self.file_name)



if __name__ == "__main__":
    # test
    doc = GenerateReport("aaa.docx")
    """doc.add_sample_info("2012-20","222","2222","DDDD","test")
    doc.add_delta_check_info({"anderson_class1":2000,"mean_class1":999})
    doc.add_graph("C://thumbnail_0.jpg")"""
    doc.add_title()
    doc.save_document("")


